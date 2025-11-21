from PyPDF2 import PdfReader
from pathlib import Path
from .config import settings
from fastapi import UploadFile
import textract
from docx import Document
from striprtf.striprtf import rtf_to_text
from typing import Optional
from telethon import TelegramClient

def process_pdf(file_path: str) -> str:
    with open(file_path, "rb") as f:
        pdf_reader = PdfReader(f)
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text()
    return text



async def save_files(files: list[UploadFile], vac_id: str) -> list[str]:
    """
    Сохранить несколько файлов резюме в папку вакансии и вернуть список путей.
    """
    dir_path = Path(settings.save_resume_path) / vac_id
    dir_path.mkdir(parents=True, exist_ok=True)

    saved_paths: list[str] = []

    for file in files:
      if not file.filename:
          # пустой инпут можно просто пропустить
          continue

      file_path = dir_path / file.filename
      content = await file.read()

    
      with open(file_path, "wb") as f:
          f.write(content)

      saved_paths.append(str(file_path))

    return saved_paths



def display_analysis(json_data):
    """
    Принимает JSON-строку или словарь Python и ВОЗВРАЩАЕТ
    структурированный отчет, содержащий Имя кандидата, "Таблицу соответствия" и "Итог".
    Если поле отсутствует, выводит '❌'.
    Автоматически удаляет маркеры блока кода ```json и ```.
    """
    import json
    
    try:
        if isinstance(json_data, str):
            # Удаляем маркеры блока кода если они есть
            json_data = json_data.strip()
            if json_data.startswith('```json'):
                json_data = json_data[7:]
            if json_data.endswith('```'):
                json_data = json_data[:-3]
            json_data = json_data.strip()
            data = json.loads(json_data)
        elif isinstance(json_data, dict):
            data = json_data
        else:
            return None
    except (json.JSONDecodeError, Exception) as e:
        return f"Ошибка парсинга JSON: {str(e)}"
    output_lines = []  # Список для хранения всех строк отчета

    # Вспомогательная функция для форматирования поля "ключ: значение"
    def format_field(key, value):
        val_str = value if value else "❌"
        return f"{key}: {val_str}"
    
    # Безопасное получение данных с проверками
    candidate = data.get("candidate", {})
    location_data = candidate.get('location', {})
    
    if isinstance(location_data, dict):
        city = location_data.get('city', None)
        country = location_data.get('country', None)
        if city == 'Нет (требуется уточнение)':
            city = None
        if country == 'Нет (требуется уточнение)':
            country = None
        if city and country:
            location = f"{city}, {country}"
        elif city:
            location = city
        elif country:
            location = country
        else:
            location = "не указано"
    else:
        location = "не указано"
    
    # --- КАНДИДАТ (только ФИО) ---
    output_lines.append("="*15 + " 👤 КАНДИДАТ " + "="*15)
    output_lines.append(format_field("ФИО", candidate.get('full_name')))
    
    # Безопасное получение даты рождения
    birth_date = candidate.get('birth_date', {})
    if isinstance(birth_date, dict):
        birth_date_str = birth_date.get('date', '❌')
    else:
        birth_date_str = '❌'
    output_lines.append(format_field("—Дата рождения", birth_date_str))
    
    # Безопасное получение зарплатных ожиданий
    summary = data.get('summary', {})
    if isinstance(summary, dict):
        salary = summary.get('salary_expectations', '❌')
    else:
        salary = '❌'
    output_lines.append(format_field("—Зарплатные ожидания", salary))
    
    output_lines.append(format_field("—Локация", location))
    
    # Безопасное получение стека технологий
    tech_stack = candidate.get('tech_stack', [])
    if isinstance(tech_stack, list) and tech_stack:
        tech_stack_str = ", ".join(tech_stack)
    else:
        tech_stack_str = "❌"
    output_lines.append(format_field("—Стек технологий", tech_stack_str))


    # --- ТАБЛИЦА СООТВЕТСТВИЯ ---
    output_lines.append("\n" + "="*12 + " ✅ ТАБЛИЦА СООТВЕТСТВИЯ " + "="*12)
    compliance = data.get("compliance_check", {})
    status_map = { "Да": "✅", "Нет (требуется уточнение)": "⚠️", "Нет (точно нет)": "❌" }
    
    must_haves = compliance.get('must_have', [])
    if must_haves and isinstance(must_haves, list):
        for req in must_haves:
            if isinstance(req, dict):
                status = req.get('status', '')
                requirement = req.get('requirement', '')
                comment = req.get('comment', '')
                
                icon = status_map.get(status, '▫️')
                if status in ["Нет (требуется уточнение)", "Нет (точно нет)"]:
                    output_lines.append(f"    {icon} {requirement}")
                    if comment:
                        clean_comment = comment.replace('⚠️', '').replace('❌', '').strip()
                        output_lines.append(f"({clean_comment})\n")
                else:
                    output_lines.append(f"    {icon} {requirement}\n")

    nice_to_haves = compliance.get('nice_to_have', [])
    if nice_to_haves and isinstance(nice_to_haves, list):
        for req in nice_to_haves:
            if isinstance(req, dict):
                status = req.get('status', '')
                requirement = req.get('requirement', '')
                comment = req.get('comment', '')
                
                icon = status_map.get(status, '▫️')
                if status in ["Нет (требуется уточнение)", "Нет (точно нет)"]:
                    output_lines.append(f"    {icon} {requirement}")
                    if comment:
                        clean_comment = comment.replace('⚠️', '').replace('❌', '').strip()
                        output_lines.append(f"({clean_comment})\n")
                else:
                    output_lines.append(f"    {icon} {requirement}\n")

    # --- ИТОГ ---
    output_lines.append("\n" + "="*17 + " 🏁 ИТОГ " + "="*17)
    if isinstance(summary, dict) and summary:
        verdict = summary.get('verdict', '❌')
        output_lines.append(format_field("Вердикт", verdict))
    else:
        output_lines.append(format_field("Вердикт", '❌'))
    output_lines.append("="*41)

    # Конвертируем в HTML для красивого отображения
    html_output = []
    for line in output_lines:
        if line.startswith("="*15 + " 👤 КАНДИДАТ"):
            html_output.append(f"<h2>👤 КАНДИДАТ</h2>")
        elif line.startswith("="*12 + " ✅ ТАБЛИЦА СООТВЕТСТВИЯ"):
            html_output.append(f"<h2>✅ ТАБЛИЦА СООТВЕТСТВИЯ</h2>")
        elif line.startswith("="*17 + " 🏁 ИТОГ"):
            html_output.append(f"<h2>🏁 ИТОГ</h2>")
        elif line.startswith("="):
            continue  # Пропускаем разделители
        elif line.strip() == "":
            html_output.append("<br>")
        elif line.startswith("    "):
            # Элементы списка
            html_output.append(f"<div style='margin-left: 20px; margin-bottom: 8px;'>{line.strip()}</div>")
        elif line.startswith("(") and line.endswith(")"):
            # Комментарии
            html_output.append(f"<div style='margin-left: 40px; color: #7f8c8d; font-style: italic; margin-bottom: 12px;'>{line}</div>")
        else:
            # Обычные поля
            if ":" in line:
                key, value = line.split(":", 1)
                html_output.append(f"<p><strong>{key.strip()}:</strong> {value.strip()}</p>")
            else:
                html_output.append(f"<p>{line}</p>")
    
    return "".join(html_output)

def norm_tg(v: str) -> str:
    v = (v or "").strip()
    return v if not v else (v if v.startswith("@") else "@" + v)


def process_doc(path: str) -> str:
    """
    Извлекает текст из .doc (старый формат Word 97–2003) с помощью textract.
    Если textract не работает (нет antiword), возвращает сообщение об ошибке.
    Возвращает очищенный текст без пустых строк.
    """
    try:
        text = textract.process(path).decode("utf-8", errors="ignore")
        lines = [line.strip() for line in text.splitlines() if line.strip()]
        return "\n".join(lines)
    except FileNotFoundError:
        print(f"⚠️ Файл не найден: {path}")
        return ""
    except textract.exceptions.ShellError as e:
        error_msg = str(e)
        if "127" in error_msg or "antiword" in error_msg.lower():
            print(f"❌ Ошибка: antiword не установлен. Пожалуйста, конвертируйте .doc файл в .docx или установите antiword")
            print(f"   Для Windows: скачайте и установите antiword с http://www.winfield.demon.nl/")
            return ""
        print(f"❌ Ошибка textract при обработке {path}: {e}")
        return ""
    except Exception as e:
        print(f"⚠️ Ошибка при чтении DOC-файла {path}: {e}")
        return ""
    


def process_docx(path: str) -> str:
    """
    Извлекает весь текст из .docx, включая таблицы и вложенные ячейки.
    Возвращает объединённый текст.
    """
    try:
        doc = Document(path)
        texts = []

        # --- Параграфы ---
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                texts.append(paragraph.text.strip())

        # --- Таблицы ---
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        texts.append(cell_text)

        # Удаляем дубликаты и объединяем
        text = "\n".join(dict.fromkeys(texts))
        return text.strip()

    except Exception as e:
        print(f"❌ Ошибка чтения DOCX: {e}")
        return ""
    

def process_rtf(path: str) -> str:
    """
    Читает RTF-файл и возвращает чистый текст.
    Работает без Pandoc.
    """
    with open(path, "r", encoding="utf-8") as f:
        content = f.read()
    text = rtf_to_text(content)
    return text

# TXT → текст
def process_txt(path: str) -> str:
    with open(path, "r", encoding="utf-8") as f:
        return f.read()
    

def parse_list(value: Optional[str]) -> list[str]:
        if not value:
            return []
        return [v.strip() for v in value.split(",") if v.strip()]


async def send_message_by_username(username: str, text: str, client: TelegramClient):
        try:
            # username можно писать без "@"
            if username.startswith("@"):
                username = username[1:]
            
            entity = await client.get_entity(username)
            await client.send_message(entity, text, parse_mode='html')
            print(f"✅ Сообщение отправлено пользователю @{username}")
            return entity
        except Exception as e:
            print(f"❌ Ошибка при отправке @{username}: {e}")
            return False