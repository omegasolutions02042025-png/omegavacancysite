# -*- coding: utf-8 -*-
"""
White Label Resume Builder:
1) build_white_label_prompt(...) — промпт для GPT/Gemini (строгие правила, Projects — обязательно).
2) generate_resume_payload_gemini(...) — запрос к Gemini, возврат JSON payload {config, content}.
3) render_resume_docx(payload) — рендер красивого .docx по JSON (Times New Roman, синие заголовки).
4) create_white_label_resume(...) — полный конвейер: кандидатский текст -> JSON -> DOCX.
5) parse_json_loose(...) — «живучий» парсер JSON из свободного текста модели.
"""

from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_BREAK
import json, re
from dotenv import load_dotenv
import os
import google.generativeai as genai
load_dotenv()
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY", "").strip()

def _extract_text_from_gemini_response(resp) -> str:
    try:
        if getattr(resp, "text", None):
            return (resp.text or "").strip()
    except Exception:
        pass
    out = []
    try:
        for cand in (getattr(resp, "candidates", None) or []):
            content = getattr(cand, "content", None)
            parts = getattr(content, "parts", None) if content else None
            if not parts:
                continue
            for p in parts:
                t = getattr(p, "text", None)
                if t: out.append(t)
    except Exception:
        pass
    return "\n".join(out).strip()

def _hex_to_rgb(hex_color: str) -> RGBColor:
    h = (hex_color or "#000000").lstrip('#')
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))

def _set_core_styles(doc: Document, font_family: str, font_size_main: int):
    style = doc.styles['Normal']
    style.font.name = font_family
    style.font.size = Pt(font_size_main)
    for sec in doc.sections:
        sec.top_margin = Inches(0.8)
        sec.bottom_margin = Inches(0.8)
        sec.left_margin = Inches(0.8)
        sec.right_margin = Inches(0.8)

def _add_section_title(doc: Document, title: str, color_hex: str, font_size_headings: int):
    p = doc.add_paragraph()
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(font_size_headings)
    r.font.color.rgb = _hex_to_rgb(color_hex)
    p.paragraph_format.line_spacing = 1.2
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)

def _add_text(doc: Document, text: str, bold: bool = False):
    if not text:
        return
    p = doc.add_paragraph(text)
    if p.runs:
        p.runs[0].bold = bool(bold)
    p.paragraph_format.line_spacing = 1.2
    p.paragraph_format.space_after = Pt(2)

def _render_skills(doc: Document, skills):
    if not skills:
        return
    if isinstance(skills, str):
        _add_text(doc, skills)
    elif isinstance(skills, list):
        _add_text(doc, ", ".join(map(str, skills)))
    elif isinstance(skills, dict):
        for k, v in skills.items():
            line = f"{k}: {', '.join(map(str, v)) if isinstance(v, list) else v}"
            _add_text(doc, line)

def _render_experience(doc: Document, exp_list: list):
    if not exp_list:
        return
    for it in exp_list:
        header = " — ".join([x for x in [it.get("company"), it.get("position")] if x])
        if header: _add_text(doc, header, bold=True)
        if it.get("period"): _add_text(doc, it["period"], bold=True)  # Сделать период полужирным
        for key in ("responsibilities", "achievements"):
            for ln in (it.get(key) or []):
                _add_text(doc, ln)
        techs = it.get("technologies") or []
        if techs:
            _add_text(doc, f"Технологии: {', '.join(map(str, techs))}")

def _render_education(doc: Document, education):
    if not education:
        return
    if isinstance(education, str):
        _add_text(doc, education); return
    for ed in education:
        line = " — ".join(filter(None, [ed.get("institution"), ed.get("degree")]))
        if line: _add_text(doc, line, bold=True)
        if ed.get("years"): _add_text(doc, ed["years"])
        if ed.get("details"): _add_text(doc, ed["details"])

def _render_projects(doc: Document, projects: list):
    if not projects:
        return
    for pr in projects:
        if pr.get("title"): _add_text(doc, pr["title"], bold=True)
        if pr.get("role"): _add_text(doc, f"Роль: {pr['role']}")
        if pr.get("period"): _add_text(doc, f"Период: {pr['period']}")
        if pr.get("description"): _add_text(doc, pr["description"])
        techs = pr.get("technologies") or []
        if techs: _add_text(doc, f"Технологии: {', '.join(map(str, techs))}")
        if pr.get("results"): _add_text(doc, f"Результаты: {pr['results']}")

def _post_fix_bold_skills(doc: Document):
    SECTION_HEADERS = {
        "РЕЗЮМЕ", "КРАТКОЕ ОПИСАНИЕ ПРОФИЛЯ",
        "КЛЮЧЕВЫЕ НАВЫКИ", "ОПЫТ РАБОТЫ", "ОБРАЗОВАНИЕ",
        "ДОПОЛНИТЕЛЬНАЯ ИНФОРМАЦИЯ", "ПРОЕКТЫ",
    }
    start_idx, end_idx = None, None
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().upper() == "КЛЮЧЕВЫЕ НАВЫКИ":
            start_idx = i + 1
            break
    if start_idx is None:
        return
    for j in range(start_idx, len(doc.paragraphs)):
        if doc.paragraphs[j].text.strip().upper() in SECTION_HEADERS:
            end_idx = j
            break
    if end_idx is None:
        end_idx = len(doc.paragraphs)
    for k in range(start_idx, end_idx):
        p = doc.paragraphs[k]
        txt = p.text
        if not txt or ":" not in txt:
            continue
        idx = txt.find(":")
        head = txt[: idx + 1].strip()
        tail = txt[idx + 1 :].lstrip()
        for r in p.runs: r.text = ""
        run_head = p.add_run(head + (" " if tail else "")); run_head.bold = True
        if tail:
            run_tail = p.add_run(tail); run_tail.bold = False

def _post_fix_inline_dicts(doc: Document):
    pattern = re.compile(r"^\s*\{\s*'title'\s*:\s*'([^']*)'\s*,\s*'items'\s*:\s*\[([^\]]*)\]\s*\}\s*$")
    for p in doc.paragraphs:
        m = pattern.match(p.text.strip())
        if not m:
            continue
        title = m.group(1).strip()
        items_raw = m.group(2).strip()
        parts = [x.strip().strip("'").strip('"') for x in items_raw.split(",") if x.strip()]
        for r in p.runs: r.text = ""
        if title:
            rt = p.add_run(title); rt.bold = True
        for it in parts:
            br = p.add_run(); br.add_break(WD_BREAK.LINE)
            p.add_run(it)

# --- Новое: форматирование даты/времени по-русски и вставка раздела «Примечание» ---

_MONTHS_RU_GEN = {
    1: "января", 2: "февраля", 3: "марта", 4: "апреля",
    5: "мая", 6: "июня", 7: "июля", 8: "августа",
    9: "сентября", 10: "октября", 11: "ноября", 12: "декабря"
}

def _format_dt_ru(dt: datetime) -> str:
    """16 августа 2025 в 11:31"""
    return f"{dt.day} {_MONTHS_RU_GEN[dt.month]} {dt.year} в {dt.strftime('%H:%M')}"

def _render_primichanie(doc: Document, color_hex: str, font_size_headings: int, utochnenie):
    """Добавляет раздел «Примечание», если utochnenie не пусто.
       Первая строка: «Резюме обновлено: <дата> (данные из оригинала).»
       Далее: по одному «Добавлено: <элемент>» для каждого пункта из utochnenie.
    """
    # Нормализуем utochnenie к списку строк
    items = []
    if isinstance(utochnenie, str):
        s = utochnenie.strip()
        if s:
            # Разрешим разделение по точке с запятой/переводам строк/запятым — но без выкидывания смысла
            candidates = re.split(r"[;\n]+", s)
            for c in candidates:
                cc = c.strip().strip("-–•").strip()
                if cc:
                    items.append(cc)
    elif isinstance(utochnenie, (list, tuple, set)):
        for x in utochnenie:
            xx = (str(x) if not isinstance(x, str) else x).strip().strip("-–•").strip()
            if xx:
                items.append(xx)

    if not items:
        return

    _add_section_title(doc, "Примечание", color_hex, font_size_headings)
    _add_text(doc, f"Резюме обновлено: {_format_dt_ru(datetime.now())} (данные из оригинала).")
    for it in items:
        _add_text(doc, f"Добавлено: {it}")

def render_resume_docx(payload: dict, vacancy_text: str = "", utochnenie=None, username = "") -> str:
    cfg = payload.get("config", {})
    cnt = payload.get("content", {})
    doc = Document()
    _set_core_styles(doc, cfg.get("font_family", "Times New Roman"), int(cfg.get("font_size_main", 12)))
    color = cfg.get("color_headings", "#1F4E79")
    hsize = int(cfg.get("font_size_headings", 14))
    sections = cfg.get("sections", [
        "ФИО","РЕЗЮМЕ","Краткое описание профиля",
        "Ключевые навыки","Опыт работы","Образование","Дополнительная информация","Проекты"
    ])
    fio = cnt.get("fio") or {}
    # Убираем отдельное отображение ФИО в начале - оно будет в секции РЕЗЮМЕ
    # Если модель убрала секцию "Дополнительная информация" (О себе),
    # но есть поле extra в content, принудительно добавим секцию,
    # чтобы пользовательский текст "О себе" не терялся.
    try:
        if cnt.get("extra"):
            upper_sections = [s.upper() for s in sections]
            if not any(s in ("ДОПОЛНИТЕЛЬНАЯ ИНФОРМАЦИЯ", "О СЕБЕ") for s in upper_sections):
                # Вставляем перед 'ПРОЕКТЫ', если она есть, иначе в конец
                try:
                    idx = upper_sections.index("ПРОЕКТЫ")
                except ValueError:
                    idx = len(sections)
                sections.insert(idx, "Дополнительная информация")
    except Exception:
        # В случае неожиданных типов оставляем поведение без изменений
        pass
    
    for sec in sections:
        su = sec.upper()
        if su == "ФИО":
            continue
        _add_section_title(doc, sec, color, hsize)
        if su == "РЕЗЮМЕ":
            if fio.get("full_name"): _add_text(doc, f"ФИО: {fio['full_name']}", bold=True)
            if cnt.get("position_grade"): _add_text(doc, f"ДОЛЖНОСТЬ: {cnt['position_grade']}", bold=True)
            if cnt.get("grade"): _add_text(doc, f"Грейд: {cnt['grade']}", bold=True)
            if fio.get("location"): _add_text(doc, f"Локация: {fio['location']}", bold=True)
            if fio.get("citizenship"): _add_text(doc, f"Гражданство: {fio['citizenship']}", bold=True)
            if fio.get("birth_date"): _add_text(doc, f"Дата рождения: {fio['birth_date']}", bold=True)
        elif su == "КРАТКОЕ ОПИСАНИЕ ПРОФИЛЯ":
            if cnt.get("summary"): _add_text(doc, cnt["summary"])
        elif su == "КЛЮЧЕВЫЕ НАВЫКИ":
            _render_skills(doc, cnt.get("skills"))
        elif su == "ОПЫТ РАБОТЫ":
            _render_experience(doc, cnt.get("experience"))
        elif su == "ОБРАЗОВАНИЕ":
            _render_education(doc, cnt.get("education"))
        elif su == "ДОПОЛНИТЕЛЬНАЯ ИНФОРМАЦИЯ":
            extra = cnt.get("extra")
            if isinstance(extra, list):
                combined_text = ", ".join(str(ln) for ln in extra if ln)
                if combined_text:
                    _add_text(doc, combined_text)
            elif isinstance(extra, str) and extra.strip():
                import re
                # убираем лишние переводы строк и двойные пробелы
                cleaned = re.sub(r"\s*\n\s*", " ", extra.strip())
                cleaned = re.sub(r"\s{2,}", " ", cleaned)
                _add_text(doc, cleaned)
        elif su == "ПРОЕКТЫ":
            _render_projects(doc, cnt.get("projects"))
    _post_fix_bold_skills(doc)
    _post_fix_inline_dicts(doc)
    
    # Выделяем технологии из вакансии жирным шрифтом
    if vacancy_text:
        technologies = _extract_technologies_from_vacancy(vacancy_text)
        _highlight_technologies_in_text(doc, technologies)

    # --- Новое: раздел «Примечание», если есть utochnenie ---
    _render_primichanie(doc, color, hsize, utochnenie)

    name_for_file = (cnt.get("fio") or {}).get("full_name") or "Name"
    date_str = datetime.now().strftime("%Y-%m-%d")
    # Use current directory instead of Linux path
    import os, re
    from urllib.parse import quote

    ILLEGAL_CHARS = re.compile(r'[\x00-\x1f<>:"/\\|?*]+')  # запреты Win/Unix

    def safe_unicode_name(s: str) -> str:
        s = (s or "").strip()
        s = s.replace(" ", "_")
        s = ILLEGAL_CHARS.sub("_", s)   # убираем слэши, двоеточия и т.п.
        s = s.strip("._-")              # края
        return s or "file"

    # dir_path: каталог, где храним WL
    # name_for_file: ФИО/название (может быть кириллица)
    # date_str: YYYY-MM-DD/ YYYYMMDD — что используешь
    # username: логин/идентификатор
    dir_path = os.path.abspath("WhiteLabel_Resume")
    os.makedirs(dir_path, exist_ok=True)
    filename = f"WhiteLabel_Resume_{safe_unicode_name(name_for_file)}_{date_str}_{safe_unicode_name(username)}.docx"
    wlfn = os.path.join(dir_path, filename)

    doc.save(wlfn)  # Unicode-имя на NTFS/EXT4 — ок

    # Отдаём фронту ссылку с URL-энкодом (иначе кириллица в href сломает маршрут)
    download_link = f"/api/wl/download/{quote(filename)}"
    return {"download_link": download_link, "filename": filename}

# --------- ДОБАВЛЕНО: устойчивое извлечение первого JSON-объекта из текста ---------
def _extract_first_json_object(s: str) -> str:
    """Извлекает первый полный JSON-объект из строки, учитывая вложенность и строки."""
    in_str = False
    esc = False
    depth = 0
    start = None
    for i, ch in enumerate(s):
        if ch == '"' and not esc:
            in_str = not in_str
        if not in_str:
            if ch == '{':
                if depth == 0:
                    start = i
                depth += 1
            elif ch == '}':
                if depth > 0:
                    depth -= 1
                if depth == 0 and start is not None:
                    return s[start:i+1]
        esc = (ch == '\\' and not esc)
    # Если JSON не закрыт, но есть начало - вернем до конца строки (для обрезанных ответов)
    if start is not None and depth > 0:
        return s[start:]
    return ""

def parse_json_loose(raw):
    if isinstance(raw, dict):
        return raw
    if not isinstance(raw, str):
        raise ValueError("Ожидалась строка с JSON.")
    
    s = raw.strip()
    print(f"DEBUG: Original response length: {len(s)}")

    # Шаг 1: Удаляем markdown блоки кода (более агрессивно)
    s = re.sub(r'^```(?:json)?\s*', '', s, flags=re.MULTILINE)
    s = re.sub(r'```\s*$', '', s, flags=re.MULTILINE)
    s = s.strip()
    
    # Удаляем возможные HTML теги или другие обертки
    s = re.sub(r'<[^>]+>', '', s)
    
    # Удаляем однострочные комментарии (// ...) и многострочные (/* ... */)
    s = re.sub(r'//.*?$', '', s, flags=re.MULTILINE)
    s = re.sub(r'/\*.*?\*/', '', s, flags=re.DOTALL)
    
    # Шаг 2: Удаляем возможные комментарии и лишний текст до/после JSON
    # Ищем первый { и последний }
    first_brace = s.find('{')
    last_brace = s.rfind('}')
    
    if first_brace != -1 and last_brace != -1 and last_brace > first_brace:
        # Извлекаем только JSON часть
        s = s[first_brace:last_brace+1]
    elif first_brace != -1:
        # Если есть открывающая скобка, но нет закрывающей - возможно JSON обрезан
        s = s[first_brace:]
    
    # Шаг 3: Пытаемся аккуратно вырезать первый законченный JSON-объект
    extracted = _extract_first_json_object(s)
    if extracted:
        try:
            result = json.loads(extracted)
            print("DEBUG: Successfully parsed using _extract_first_json_object")
            return result
        except json.JSONDecodeError as e:
            print(f"DEBUG: _extract_first_json_object parse failed: {e}")
            # Если извлеченный JSON неполный, попробуем его восстановить
            if extracted.endswith('"') or extracted.endswith(','):
                # Возможно, JSON обрезан
                pass
    
    # Шаг 4: Проверяем на обрезанный JSON
    open_braces = s.count('{')
    close_braces = s.count('}')
    open_brackets = s.count('[')
    close_brackets = s.count(']')
    
    if open_braces > close_braces or open_brackets > close_brackets:
        print(f"WARNING: JSON appears to be truncated. Open braces: {open_braces}, Close braces: {close_braces}, Open brackets: {open_brackets}, Close brackets: {close_brackets}")
    
    # Шаг 5: Пробуем парсить как есть
    try:
        return json.loads(s)
    except json.JSONDecodeError as e:
        print(f"DEBUG: First parse attempt failed: {e}")
    
    # Шаг 6: Пробуем извлечь JSON из фигурных скобок
    try:
        start, end = s.find("{"), s.rfind("}")
        if start != -1 and end != -1 and end > start:
            json_part = s[start:end+1]
            print(f"DEBUG: Extracted JSON part length: {len(json_part)}")
            return json.loads(json_part)
    except json.JSONDecodeError as e:
        print(f"DEBUG: Second parse attempt failed: {e}")
    
    # Шаг 7: Удаляем невидимые символы, управляющие символы и пробуем еще раз
    # Удаляем невидимые символы Unicode
    s2 = re.sub(r"[\u200b-\u200f\u202a-\u202e\u00a0]", "", s)
    s2 = s2.replace('\ufeff', '')  # BOM
    # Удаляем управляющие символы (кроме \n, \r, \t которые могут быть в строках)
    s2 = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f]', '', s2)
    
    try:
        return json.loads(s2)
    except json.JSONDecodeError as e:
        print(f"DEBUG: Parse attempt with cleaned string failed: {e}")
    
    # Шаг 8: Пробуем восстановить JSON с исправлением распространенных ошибок
    s3 = s2
    
    # Исправляем незакрытые строки и неэкранированные символы новой строки в строках
    in_string = False
    escape_next = False
    result_chars = []
    for i, char in enumerate(s3):
        if escape_next:
            result_chars.append(char)
            escape_next = False
            continue
        if char == '\\':
            result_chars.append(char)
            escape_next = True
            continue
        if char == '"':
            in_string = not in_string
            result_chars.append(char)
        elif char == '\n' and in_string:
            # Экранируем неэкранированные символы новой строки внутри строк
            result_chars.append('\\n')
        elif char == '\r' and in_string:
            # Экранируем неэкранированные символы возврата каретки внутри строк
            result_chars.append('\\r')
        elif char == '\t' and in_string:
            # Экранируем неэкранированные табуляции внутри строк
            result_chars.append('\\t')
        else:
            result_chars.append(char)
    
    # Если строка не закрыта, закрываем её
    if in_string:
        print("DEBUG: Found unterminated string, closing it")
        s3 = ''.join(result_chars) + '"'
    else:
        s3 = ''.join(result_chars)
    
    # Исправляем одинарные кавычки на двойные (осторожно, чтобы не сломать экранированные)
    # Сначала заменяем только вне строк
    s3_temp = s3
    in_str = False
    esc = False
    fixed_chars = []
    for char in s3_temp:
        if esc:
            fixed_chars.append(char)
            esc = False
            continue
        if char == '\\':
            fixed_chars.append(char)
            esc = True
            continue
        if char == '"':
            in_str = not in_str
            fixed_chars.append(char)
        elif char == "'" and not in_str:
            fixed_chars.append('"')
        else:
            fixed_chars.append(char)
    s3 = ''.join(fixed_chars)
    
    # Удаляем запятые перед закрывающими скобками/фигурными скобками
    s3 = re.sub(r',(\s*[}\]])', r'\1', s3)
    
    # Исправляем незакавыченные ключи: ключ без кавычек перед : должен быть в кавычках
    s3 = re.sub(r'(?<=[{\s,])([a-zA-Z_][a-zA-Z0-9_]*)(?=\s*:)', r'"\1"', s3)
    
    # Исправляем отсутствующие двоеточия: если есть "key" перед значением без двоеточия, добавляем
    s3 = re.sub(r'"([^"]*?)"\s+(?=["\[{0-9true false])', r'"\1": ', s3)
    
    # Исправляем невалидные escape-последовательности (например, \x вместо \\x)
    # Но только вне строк, чтобы не сломать валидные экранированные символы
    # Это сложно сделать правильно, поэтому пропускаем этот шаг
    
    # Исправляем случаи, когда значение обрезано в середине (например, "text без закрывающей кавычки)
    # Находим незакрытые строки и закрываем их перед закрывающими скобками
    if s3.count('"') % 2 != 0:
        # Нечетное количество кавычек - есть незакрытая строка
        print("DEBUG: Detected unclosed string, attempting to fix")
        # Находим последнюю открывающую кавычку без закрывающей
        last_open_quote = s3.rfind('"')
        if last_open_quote > 0:
            # Проверяем, не является ли это закрывающей кавычкой
            before_quote = s3[:last_open_quote]
            quote_count_before = before_quote.count('"')
            if quote_count_before % 2 == 0:
                # Это открывающая кавычка без закрывающей
                # Закрываем её перед последней закрывающей скобкой
                last_brace = s3.rfind('}')
                if last_brace > last_open_quote:
                    s3 = s3[:last_brace] + '"' + s3[last_brace:]
    
    # Удаляем запятые в конце объектов/массивов (повторно, на случай если что-то пропустили)
    s3 = re.sub(r',(\s*[}\]])', r'\1', s3)
    
    # Пробуем завершить обрезанный JSON, добавляя недостающие закрывающие скобки
    open_braces = s3.count('{')
    close_braces = s3.count('}')
    open_brackets = s3.count('[')
    close_brackets = s3.count(']')
    
    if open_braces > close_braces:
        missing_braces = open_braces - close_braces
        print(f"DEBUG: Attempting to complete truncated JSON by adding {missing_braces} closing braces")
        # Удаляем возможные запятые и пробелы в конце перед добавлением скобок
        s3 = s3.rstrip(',\n\r\t ') + '}' * missing_braces
    
    if open_brackets > close_brackets:
        missing_brackets = open_brackets - close_brackets
        print(f"DEBUG: Attempting to complete truncated arrays by adding {missing_brackets} closing brackets")
        s3 = s3.rstrip(',\n\r\t ') + ']' * missing_brackets
    
    try:
        return json.loads(s3)
    except json.JSONDecodeError as e2:
        print(f"DEBUG: Even after fixes, parsing failed: {e2}")
        error_pos = getattr(e2, 'pos', 0)
        if error_pos > 0:
            start_pos = max(0, error_pos - 200)
            end_pos = min(len(s3), error_pos + 200)
            print(f"DEBUG: Problematic JSON around error position {error_pos}:")
            print(f"DEBUG: Context: {repr(s3[start_pos:end_pos])}")
        
        # Пробуем еще более агрессивное исправление: обрезаем до последней закрывающей скобки
        try:
            last_brace = s3.rfind('}')
            if last_brace > 0:
                s4 = s3[:last_brace+1]
                result = json.loads(s4)
                print("DEBUG: Successfully parsed by truncating to last closing brace")
                return result
        except json.JSONDecodeError as e3:
            print(f"DEBUG: Truncation to last brace failed: {e3}")
        
        # Финальная попытка: закрываем все незакрытые скобки и фигурные скобки
        print("DEBUG: Attempting final fix: closing all unclosed brackets")
        s5 = s3
        try:
            # Подсчитываем незакрытые скобки и фигурные скобки
            brace_depth = 0
            bracket_depth = 0
            in_str = False
            esc = False
            for char in s5:
                if esc:
                    esc = False
                    continue
                if char == '\\':
                    esc = True
                    continue
                if char == '"':
                    in_str = not in_str
                    continue
                if not in_str:
                    if char == '{':
                        brace_depth += 1
                    elif char == '}':
                        brace_depth = max(0, brace_depth - 1)
                    elif char == '[':
                        bracket_depth += 1
                    elif char == ']':
                        bracket_depth = max(0, bracket_depth - 1)
            
            # Закрываем все незакрытые скобки и фигурные скобки
            s5 = s5.rstrip(',\n\r\t ')
            s5 += ']' * bracket_depth + '}' * brace_depth
            result = json.loads(s5)
            print(f"DEBUG: Successfully parsed after closing {bracket_depth} brackets and {brace_depth} braces")
            return result
        except json.JSONDecodeError as e4:
            print(f"DEBUG: Final fix also failed: {e4}")
            
            # Сохраняем проблемный JSON в файл для анализа
            try:
                with open("debug_json_error.txt", "w", encoding="utf-8") as f:
                    f.write(f"First parse error: {e if 'e' in locals() else 'N/A'}\n")
                    f.write(f"Fixed error (e2): {e2}\n")
                    f.write(f"Final error (e4): {e4}\n")
                    f.write(f"Error position: {error_pos if 'error_pos' in locals() else 'N/A'}\n")
                    f.write(f"Original response length: {len(s)}\n")
                    f.write(f"Cleaned string length: {len(s2)}\n")
                    f.write(f"Fixed string length: {len(s3)}\n")
                    f.write(f"\nOriginal problematic JSON (first 5000 chars):\n{s2[:5000]}\n\n")
                    f.write(f"After fixes (first 5000 chars):\n{s3[:5000]}\n\n")
                    if 's5' in locals() and s5:
                        f.write(f"Attempted final fix (first 5000 chars):\n{s5[:5000]}")
            except Exception as save_err:
                print(f"DEBUG: Failed to save debug file: {save_err}")
            
            raise ValueError(f"Не удалось распарсить JSON после всех попыток. Последняя ошибка: {e4}. Детали сохранены в debug_json_error.txt. Возможно, ответ от модели был обрезан или содержит фундаментальные синтаксические ошибки.")

def _extract_technologies_from_vacancy(vacancy_text: str) -> list:
    """Извлекает технологии и ключевые слова из текста вакансии"""
    import re
    
    # Общие технологии и фреймворки
    tech_patterns = [
        r'\b(?:Python|Java|JavaScript|TypeScript|C\#|C\+\+|PHP|Ruby|Go|Rust|Kotlin|Swift)\b',
        r'\b(?:React|Angular|Vue|Django|Flask|Spring|Laravel|Express|Node\.js)\b',
        r'\b(?:PostgreSQL|MySQL|MongoDB|Redis|Elasticsearch|Oracle|SQL Server)\b',
        r'\b(?:Docker|Kubernetes|AWS|Azure|GCP|Jenkins|GitLab|GitHub)\b',
        r'\b(?:Linux|Windows|MacOS|Ubuntu|CentOS)\b',
        r'\b(?:Git|SVN|Mercurial)\b',
        r'\b(?:REST|GraphQL|API|JSON|XML|SOAP)\b',
        r'\b(?:HTML|CSS|SASS|LESS|Bootstrap|Tailwind)\b',
        r'\b(?:Webpack|Vite|Babel|ESLint|Prettier)\b',
        r'\b(?:Terraform|Ansible|Puppet|Chef)\b',
        r'\b(?:Prometheus|Grafana|ELK|Splunk)\b',
        r'\b(?:Kafka|RabbitMQ|ActiveMQ)\b',
        r'\b(?:Hadoop|Spark|Airflow|Greenplum)\b'
    ]
    
    technologies = set()
    for pattern in tech_patterns:
        matches = re.findall(pattern, vacancy_text, re.IGNORECASE)
        technologies.update(matches)
    
    return list(technologies)

def _highlight_technologies_in_text(doc: Document, technologies: list):
    """Выделяет технологии жирным шрифтом в документе"""
    if not technologies:
        return
        
    import re
    
    for paragraph in doc.paragraphs:
        if not paragraph.text.strip():
            continue
            
        original_text = paragraph.text
        has_matches = False
        
        # Проверяем, есть ли технологии в этом параграфе
        for tech in technologies:
            if re.search(r'\b' + re.escape(tech) + r'\b', original_text, re.IGNORECASE):
                has_matches = True
                break
        
        if not has_matches:
            continue
            
        # Очищаем параграф и пересоздаем с выделением
        paragraph.clear()
        
        remaining_text = original_text
        while remaining_text:
            # Находим ближайшее совпадение
            earliest_match = None
            earliest_pos = len(remaining_text)
            matched_tech = None
            
            for tech in technologies:
                match = re.search(r'\b' + re.escape(tech) + r'\b', remaining_text, re.IGNORECASE)
                if match and match.start() < earliest_pos:
                    earliest_pos = match.start()
                    earliest_match = match
                    matched_tech = tech
            
            if earliest_match is None:
                # Нет больше совпадений, добавляем остальной текст
                paragraph.add_run(remaining_text)
                break
            
            # Добавляем текст до совпадения
            if earliest_pos > 0:
                paragraph.add_run(remaining_text[:earliest_pos])
            
            # Добавляем совпадение жирным
            bold_run = paragraph.add_run(earliest_match.group())
            bold_run.bold = True
            
            # Продолжаем с оставшимся текстом
            remaining_text = remaining_text[earliest_match.end():]

def build_prompt_simple(candidate_text: str, vacancy_text: str) -> str:
    return f"""
Верни СТРОГО валидный JSON без какого-либо текста вне объекта. Никаких комментариев, кода, маркдауна, пояснений.
Только объект {{"config":{{...}}, "content":{{...}}}}.
White Label: не включай контакты и email. Сохрани ВСЁ содержание без сокращений.
Если нет Summary — создай 3–5 предложений. Определи должность по вакансии.
ГРЕЙД: определи только как Senior, Middle или Junior на основе опыта работы.
ГРАЖДАНСТВО: определи из локации и укажи как РФ (для России/Москвы), РБ (для Беларуси/Минска), или возьми из резюме если указано.
ПРОЕКТЫ обязательны: найди все даже если они спрятаны в обязанностях/Обо мне.
ДОПОЛНИТЕЛЬНАЯ ИНФОРМАЦИЯ (О себе): форматируй компактно, объединяй характеристики в связный текст без лишних переносов строк.

⚠️ КРИТИЧЕСКОЕ ТРЕБОВАНИЕ К JSON (ПРОВЕРЯЙ ПЕРЕД ВЫВОДОМ):
1. ТОЛЬКО один JSON-объект: {{ ... }}
2. НЕ добавляй текст до/после JSON
3. Все ключи в двойных кавычках: "key"
4. Все строки в двойных кавычках: "value"
5. КАЖДЫЙ ключ ОБЯЗАТЕЛЬНО имеет ":  " (двоеточие с пробелом). Правильно: "key": "value". НЕПРАВИЛЬНО: "key" "value"
6. Если кавычка в значении — экранируй: \"
7. Нет запятых после последнего элемента перед }} или ]]
8. Посчитай: {{ должны равняться }}, [ должны равняться ]
9. ЕСЛИ ЧТО-ТО НЕ ТАК — ПЕРЕДЕЛАЙ JSON полностью ПЕРЕД ВЫВОДОМ

ЗАПОМНИ: каждому ключу НЕОБХОДИМО двоеточие сразу после кавычки "key": 
Неправильные примеры (ЗАПРЕЩЕНЫ):
- "key" value  (нет двоеточия)
- "key"value   (нет двоеточия)
- "key" "value" (нет двоеточия)
Правильный пример:
- "key": "value"  ✓
- "key": 123  ✓
- "key": true  ✓
- "key": ["item1", "item2"]  ✓

Схема:
{{
 "config": {{
   "output_format": "docx",
   "font_family": "Times New Roman",
   "font_size_main": 12,
   "font_size_headings": 14,
   "color_headings": "#1F4E79",
   "language": "ru",
   "sections": [
     "ФИО","РЕЗЮМЕ","Краткое описание профиля",
     "Ключевые навыки","Опыт работы","Образование","Дополнительная информация","Проекты"
   ],
   "white_label": true,
   "exclude_contacts": true,
   "exclude_email": true
 }},
 "content": {{
   "fio": {{"full_name":"","location":"","citizenship":"","birth_date":""}},
   "position_grade":"", "grade":"", "summary":"",
   "skills": {{}},
   "experience":[{{"company":"","position":"","period":"","responsibilities":[],"technologies":[],"achievements":[]}}],
   "education":[{{"institution":"","degree":"","years":"","details":""}}],
   "extra": "текст О себе одним абзацем или список коротких фраз",
   "projects":[{{"title":"","role":"","period":"","description":"","technologies":[],"results":""}}]
 }}
}}

ИСХОДНОЕ РЕЗЮМЕ:
{candidate_text}

ТРЕБОВАНИЯ ВАКАНСИИ:
{vacancy_text}
"""

async def generate_payload_once(
                          candidate_text: str,
                          vacancy_text: str,
                          model : genai.GenerativeModel) -> dict:
    
    prompt = build_prompt_simple(candidate_text, vacancy_text)
    resp = await model.generate_content_async(
        prompt,
        generation_config=genai.types.GenerationConfig(
            temperature=0.1,
            max_output_tokens=20480,  # Максимум для Gemini-2.5 (было 12288)
            response_mime_type="application/json"  # <-- Жёсткий JSON-режим
        )
    )
    raw = _extract_text_from_gemini_response(resp)
    
    # Проверяем finish_reason
    finish_reason = None
    try:
        if resp.candidates:
            finish_reason = getattr(resp.candidates[0], "finish_reason", None)
    except Exception:
        pass
    
    # finish_reason=2 означает MAX_TOKENS — ответ был обрезан
    if finish_reason == 2:
        print(f"WARNING: Ответ от Gemini был обрезан (MAX_TOKENS). Попытаемся парсить частичный JSON.")
        if not raw:
            raise ValueError(f"Ответ от Gemini был обрезан (MAX_TOKENS) и пуст. Увеличьте max_output_tokens или сократите входные данные.")
        # Продолжим парсинг, даже если ответ обрезан
    elif not raw:
        safety = None
        try:
            if resp.candidates:
                safety = getattr(resp.candidates[0], "safety_ratings", None)
        except Exception:
            pass
        raise ValueError(f"Пустой ответ от Gemini. finish_reason={finish_reason}, safety={safety}")
    data = parse_json_loose(raw)
    if not isinstance(data, dict) or "config" not in data or "content" not in data:
        raise ValueError("Модель не вернула JSON с ключами {config, content}.")
    cfg = data.setdefault("config", {})
    if "sections" in cfg and "Проекты" not in cfg["sections"]:
        cfg["sections"].append("Проекты")
    return data

async def create_white_label_resume(candidate_text: str,
                                   vacancy_text: str,
                                   model : genai.GenerativeModel,
                                   utochnenie=None,
                                   username = ""):
    
    payload = await generate_payload_once(candidate_text, vacancy_text, model)
    filename = render_resume_docx(payload, vacancy_text, utochnenie=utochnenie, username = username)
    return filename

#===== Пример использования (раскомментируй, подставь API ключ и тексты) =====

import os
from dotenv import load_dotenv
load_dotenv()
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
if __name__ == "__main__":
    candidate = '''
      
Contact
 Apt 321, 31, Veshnyakovskaya
 str., Moscow, 111538, Russian
 Federation.
 7-916-685-9607 (Mobile)
 sergts@mail.ru
 www.linkedin.com/in/sergts
 (LinkedIn)
 github.com/itboss2 (Portfolio)
 Top Skills
 System Architects
 Telecommunications Billing
 WebSphere ESB
 Languages
 English - C1 Certified (Full
 Professional)
 Bulgarian (Elementary)
 Russian - C2 Certifed (Native or
 Bilingual)
 Byelorussian - C2 Certified (Full
 Professional)
 Polish (Limited Working)
 German (Elementary)
 Ukrainian (Elementary)
 Certifications
 Palo Alto Networks Accredited
 Configuration Engineer (ACE) 
PAN-OS 6.1
 Cisco Network Support Engineer
 VMware Technical Sales
 Professional 5 for 6 competencies:
 Infrastructure Virtualization, Desktop
 Virtualization, Business Continuity,
 Virtualization of Business Critical
 Applications, Infrastructure as a
 Service, Management
 IBM Certified Database Associate 
DB2 10.1 Fundamentals
 Honors-Awards
 3rd prize winner of Russia Huawei
 2017 Channel Partner Skill
 Competition
 Sergey Tsuprikov, Tech
 Advisor, IBM DS, LSSBB,
 MCP,Mentor
 IT architect and project manager with 10+ years of experience in
 Big Data, Data Science, enterprise software design area | Banking
 | Telecom | FinTech | Data Lake | DWH | BI | Data Driven | Data
 Management | Data Quality
 Moscow, Moscow City, Russia
 Summary
 12+ years of experience as an IT architect for enterprise software
 (like OEBSl) design, complicated cross-platform data migration and
 integration, distributed data centers turn-key design.
 7+ years of experience as a project manager for projects up to 15M
 USD. 
7+ years of experience as a business and/or system analyst, mostly
 in the banking area. 
5+ years of experience as an out-staff instructor (PMI PMBOK,
 PRINCE2, ITIL, IBM solutions).
 10+ years of experience as a subject matter expert (SME) in banking
 and finance, sales and marketing, logistics, manufacturing, retail,
 FMCG. 
5+ years of successful daily work face-to-face and remote with
 Indian banking analysts from Oracle Corp. 
My mobile phone: 7-916-685-9607 (9:00-22:00 GMT+3, Moscow),
 Telegram: @sergts1 .
 My private email - sergts@mail.ru (please, use bossit@gmail.com
 only in case of troubles with sending to sergts@mail.ru).
 I have 30+ years of experience in IT. Have received 120+ worldwide
 IT certificates after exams (i.e. 70+ technical from IBM): Program
 & Portfolio Management Expert (#35), Project Management
 Expert (PME) #000412, IBM Data Science Professional, VMware
 Certified Professional 4 (#63533), EMCTAe, Six Sigma Black Belt
 Professional, MCP, etc. 
 Page 1 of 9
  ...
'''
    vacancy = '''
    BD-10128 (https://t.me/omega_vacancy_bot?start=3093_BD-10128)
📅 Дата публикации: 08.10.2025 12:13
... (укорочено для примера) ...
'''
    # # Пример уточнений:
    # utochnenie = [
    #     "владение французским языком (уровень не указан)",
    #     "опыт с Greenplum (подтвердить стек)",
    #     "сертификация AWS (подтвердить год)"
    # ]
    # fn = create_white_label_resume_once(GEMINI_API_KEY, candidate, vacancy, utochnenie=utochnenie)
    # print("Готово:", fn)
